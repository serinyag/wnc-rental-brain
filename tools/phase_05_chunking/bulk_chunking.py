from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook

from .chunking import (
    CHUNKING_STRATEGY_CODE,
    CHUNKING_STRATEGY_VERSION,
    MAX_CHUNK_TOKENS,
    ChunkRecord,
    PilotDocumentConfig,
    PilotDocumentResult,
    _build_chunk,
    _cell_text,
    _checklist_question_label,
    _extract_heading_level,
    _extract_top_level_number,
    _is_checklist_subheading,
    _nonempty_paragraphs,
    _split_semantic_section,
    _strip_numeric_prefix,
    _with_ordinal,
    generate_document_result,
    pilot_document_map,
    repo_root,
)


@dataclass(frozen=True)
class BulkDocumentConfig:
    document_code: str
    canonical_title: str
    relative_path: str
    parser_kind: str
    parser_version: str
    top_level_include_numbers: tuple[int, ...] = ()
    sheet_names: tuple[str, ...] = ()


BULK_DOCUMENTS: tuple[BulkDocumentConfig, ...] = (
    BulkDocumentConfig(
        document_code="CF-003",
        canonical_title="Studio Rental Terms",
        relative_path="sources/phase-01-03/Client Facing Docs/Studio Space _ Terms and Conditions.docx",
        parser_kind="docx_heading_outline",
        parser_version="docx_heading_outline_v1",
    ),
    BulkDocumentConfig(
        document_code="CF-005",
        canonical_title="Full Venue Rental Terms",
        relative_path="sources/phase-01-03/Client Facing Docs/Full Venue _ Rental Terms and Conditions.docx",
        parser_kind="docx_heading_outline",
        parser_version="docx_heading_outline_v1",
    ),
    BulkDocumentConfig(
        document_code="CF-007",
        canonical_title="WNC Rental Agreement Template",
        relative_path="sources/phase-01-03/Client Facing Docs/WNC Rental Agreement Template.docx",
        parser_kind="docx_heading_outline_with_tables",
        parser_version="docx_heading_outline_with_tables_v1",
    ),
    BulkDocumentConfig(
        document_code="GOV-001",
        canonical_title="WNC Rental Knowledge Inventory",
        relative_path="sources/phase-01-03/Knowledge Governance/WNC Rental Knowledge Inventory.xlsm",
        parser_kind="xlsx_governance_inventory",
        parser_version="xlsx_governance_inventory_v1",
    ),
    BulkDocumentConfig(
        document_code="GOV-002",
        canonical_title="WNC Rental Policy Decisions & Change Log",
        relative_path="sources/phase-01-03/Knowledge Governance/WNC Rental Policy Decisions & Change Log.xlsm",
        parser_kind="xlsx_policy_decision_log",
        parser_version="xlsx_policy_decision_log_v1",
    ),
    BulkDocumentConfig(
        document_code="OPS-001",
        canonical_title="WNC Venue Rental Operations Manual",
        relative_path="sources/phase-01-03/Venue & Operations/WNC Venue Rental Operations Manual.docx",
        parser_kind="docx_heading_outline",
        parser_version="docx_heading_outline_v1",
    ),
    BulkDocumentConfig(
        document_code="OPS-002",
        canonical_title="WNC Venue Technical & Equipment Inventory",
        relative_path="sources/phase-01-03/Venue & Operations/WNC Venue Technical & Equipment Inventory.xlsm",
        parser_kind="xlsx_technical_inventory",
        parser_version="xlsx_technical_inventory_v1",
    ),
    BulkDocumentConfig(
        document_code="OPS-003",
        canonical_title="WNC Capacity & Space Use Rules",
        relative_path="sources/phase-01-03/Venue & Operations/WNC Venue Technical & Equipment Inventory.xlsm",
        parser_kind="xlsx_capacity_space_rules",
        parser_version="xlsx_capacity_space_rules_v1",
    ),
    BulkDocumentConfig(
        document_code="SERV-001",
        canonical_title="WNC Rental Services Catalogue",
        relative_path="sources/phase-01-03/Catalogues/WNC Rental Services Catalogue.xlsm",
        parser_kind="xlsx_service_catalogue",
        parser_version="xlsx_service_catalogue_v2",
    ),
    BulkDocumentConfig(
        document_code="SERV-003",
        canonical_title="WNC Catering, Beverage & Supplier Catalogue",
        relative_path="sources/phase-01-03/Catalogues/WNC Catering, Beverage & Supplier Catalogue.xlsm",
        parser_kind="xlsx_catering_supplier_catalogue",
        parser_version="xlsx_catering_supplier_catalogue_v1",
    ),
    BulkDocumentConfig(
        document_code="SERV-004",
        canonical_title="External Supplier Requirements",
        relative_path="sources/phase-01-03/Catalogues/WNC Catering, Beverage & Supplier Catalogue.xlsm",
        parser_kind="xlsx_external_supplier_requirements",
        parser_version="xlsx_external_supplier_requirements_v1",
    ),
    BulkDocumentConfig(
        document_code="TPL-001",
        canonical_title="Studio Rental Proposal Template",
        relative_path="sources/phase-01-03/Checklists + Templates/Proposal Templates/Studio Rental Proposal Template.docx",
        parser_kind="docx_numbered_template_sections",
        parser_version="docx_numbered_template_sections_v1",
    ),
    BulkDocumentConfig(
        document_code="TPL-002",
        canonical_title="Entire Venue Proposal Template",
        relative_path="sources/phase-01-03/Checklists + Templates/Proposal Templates/Entire Venue Proposal Template.docx",
        parser_kind="docx_numbered_template_sections",
        parser_version="docx_numbered_template_sections_v1",
    ),
    BulkDocumentConfig(
        document_code="TPL-003",
        canonical_title="Custom Scope Proposal Template",
        relative_path="sources/phase-01-03/Checklists + Templates/Proposal Templates/Custom Scope Proposal Template.docx",
        parser_kind="docx_numbered_template_sections",
        parser_version="docx_numbered_template_sections_v1",
    ),
    BulkDocumentConfig(
        document_code="TPL-004",
        canonical_title="Production Coordination Proposal Template",
        relative_path="sources/phase-01-03/Checklists + Templates/Proposal Templates/Production Coordination Proposal Template.docx",
        parser_kind="docx_numbered_template_sections",
        parser_version="docx_numbered_template_sections_v1",
    ),
    BulkDocumentConfig(
        document_code="TPL-005",
        canonical_title="Full Production Proposal Template",
        relative_path="sources/phase-01-03/Checklists + Templates/Proposal Templates/Full Production Proposal Template.docx",
        parser_kind="docx_numbered_template_sections",
        parser_version="docx_numbered_template_sections_v1",
    ),
    BulkDocumentConfig(
        document_code="TPL-006",
        canonical_title="WNC Rental Email Template Library",
        relative_path="sources/phase-01-03/Checklists + Templates/WNC Rental Email Template Library.docx",
        parser_kind="docx_template_library",
        parser_version="docx_template_library_v2",
    ),
    BulkDocumentConfig(
        document_code="TPL-007",
        canonical_title="Discovery Call Checklist",
        relative_path="sources/phase-01-03/Checklists + Templates/WNC Rental Discovery Call & Site Visit Checklist.docx",
        parser_kind="docx_checklist_sections",
        parser_version="docx_checklist_sections_v1",
        top_level_include_numbers=(1, 3),
    ),
    BulkDocumentConfig(
        document_code="TPL-008",
        canonical_title="Site Visit Checklist",
        relative_path="sources/phase-01-03/Checklists + Templates/WNC Rental Discovery Call & Site Visit Checklist.docx",
        parser_kind="docx_numbered_checklist_sections",
        parser_version="docx_numbered_checklist_sections_v1",
        top_level_include_numbers=(2,),
    ),
    BulkDocumentConfig(
        document_code="TPL-009",
        canonical_title="Event Handover Checklist",
        relative_path="sources/phase-01-03/Checklists + Templates/WNC Rental Event Handover & Final Readiness Checklist.docx",
        parser_kind="docx_numbered_checklist_sections",
        parser_version="docx_numbered_checklist_sections_v1",
        top_level_include_numbers=(1, 2, 3, 4, 6, 7, 8),
    ),
    BulkDocumentConfig(
        document_code="TPL-010",
        canonical_title="Final Readiness Checklist",
        relative_path="sources/phase-01-03/Checklists + Templates/WNC Rental Event Handover & Final Readiness Checklist.docx",
        parser_kind="docx_numbered_checklist_sections",
        parser_version="docx_numbered_checklist_sections_v1",
        top_level_include_numbers=(5,),
    ),
    BulkDocumentConfig(
        document_code="TPL-013",
        canonical_title="Rental Close-Out Checklist",
        relative_path="sources/phase-01-03/Checklists + Templates/WNC Rental Close-Out Checklist.docx",
        parser_kind="docx_linear_checklist_sections",
        parser_version="docx_linear_checklist_sections_v1",
    ),
)


def bulk_document_map() -> dict[str, BulkDocumentConfig]:
    return {config.document_code: config for config in BULK_DOCUMENTS}


def generate_bulk_results(document_codes: Iterable[str] | None = None) -> list[PilotDocumentResult]:
    configs = BULK_DOCUMENTS
    if document_codes is not None:
        allowed = set(document_codes)
        configs = tuple(config for config in BULK_DOCUMENTS if config.document_code in allowed)
    return [generate_bulk_document_result(config) for config in configs]


def generate_bulk_document_result(config: BulkDocumentConfig) -> PilotDocumentResult:
    pilot_config = pilot_document_map().get(config.document_code)
    if pilot_config is not None:
        return generate_document_result(pilot_config)

    if config.parser_kind in {
        "docx_heading_outline",
        "docx_template_library",
        "docx_checklist_sections",
        "xlsx_service_catalogue",
    }:
        return generate_document_result(
            PilotDocumentConfig(
                document_code=config.document_code,
                canonical_title=config.canonical_title,
                relative_path=config.relative_path,
                parser_kind=config.parser_kind,
                parser_version=config.parser_version,
                top_level_include_numbers=config.top_level_include_numbers,
            )
        )

    path = repo_root() / config.relative_path
    if config.parser_kind == "docx_heading_outline_with_tables":
        chunks = _chunk_docx_heading_outline_with_tables(path, config.canonical_title)
    elif config.parser_kind == "docx_numbered_template_sections":
        chunks = _chunk_docx_numbered_sections(
            path,
            config.canonical_title,
            config.top_level_include_numbers,
            intro_heading="Proposal header and working fields",
            question_prefix="How should the proposal section",
        )
    elif config.parser_kind == "docx_numbered_checklist_sections":
        chunks = _chunk_docx_numbered_sections(
            path,
            config.canonical_title,
            config.top_level_include_numbers,
            intro_heading=None,
            question_prefix="What should be captured in",
        )
    elif config.parser_kind == "docx_linear_checklist_sections":
        chunks = _chunk_docx_linear_checklist_sections(path, config.canonical_title)
    elif config.parser_kind == "xlsx_governance_inventory":
        chunks = _chunk_governance_inventory(path, config.canonical_title)
    elif config.parser_kind == "xlsx_policy_decision_log":
        chunks = _chunk_policy_decision_log(path, config.canonical_title)
    elif config.parser_kind == "xlsx_technical_inventory":
        chunks = _chunk_technical_inventory(path, config.canonical_title)
    elif config.parser_kind == "xlsx_capacity_space_rules":
        chunks = _chunk_capacity_space_rules(path, config.canonical_title)
    elif config.parser_kind == "xlsx_catering_supplier_catalogue":
        chunks = _chunk_catering_supplier_catalogue(path, config.canonical_title)
    elif config.parser_kind == "xlsx_external_supplier_requirements":
        chunks = _chunk_external_supplier_requirements(path, config.canonical_title)
    else:
        raise ValueError(f"Unsupported bulk parser_kind: {config.parser_kind}")

    return PilotDocumentResult(
        document_code=config.document_code,
        canonical_title=config.canonical_title,
        relative_path=config.relative_path,
        parser_kind=config.parser_kind,
        parser_version=config.parser_version,
        chunking_strategy_code=CHUNKING_STRATEGY_CODE,
        chunking_strategy_version=CHUNKING_STRATEGY_VERSION,
        max_chunk_tokens=MAX_CHUNK_TOKENS,
        chunks=chunks,
    )


def _iter_docx_blocks(document: Document) -> Iterable[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _table_to_lines(table: Table) -> list[str]:
    rows: list[list[str]] = []
    for row in table.rows:
        values = [" ".join(cell.text.split()) for cell in row.cells]
        if any(values):
            rows.append(values)
    if not rows:
        return []

    if len(rows[0]) == 2 and all(len(row) == 2 for row in rows):
        return [f"{row[0]}: {row[1]}" for row in rows if row[0] or row[1]]

    headers = rows[0]
    lines: list[str] = []
    for row in rows[1:]:
        parts = [
            f"{header}: {value}"
            for header, value in zip(headers, row)
            if header and value
        ]
        if parts:
            lines.append("; ".join(parts))
        else:
            lines.append(" | ".join(value for value in row if value))
    return lines


def _chunk_docx_heading_outline_with_tables(path: Path, title: str) -> list[ChunkRecord]:
    document = Document(path)
    heading_stack: list[tuple[int, str]] = []
    intro_body: list[str] = []
    current_heading: str | None = None
    current_path: str | None = None
    current_body: list[str] = []
    sections: list[tuple[str, str, list[str]]] = []

    def flush_current() -> None:
        nonlocal current_body
        if current_heading and current_body:
            sections.append((current_heading, current_path or current_heading, current_body[:]))
        current_body = []

    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = " ".join(block.text.split())
            if not text:
                continue
            level = _extract_heading_level(block.style.name if block.style else "Normal")
            if level is not None:
                flush_current()
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, text))
                current_heading = text
                current_path = " > ".join(item[1] for item in heading_stack)
                continue
            if current_heading is None:
                intro_body.append(text)
            else:
                current_body.append(text)
            continue

        table_lines = _table_to_lines(block)
        if current_heading is None:
            intro_body.extend(table_lines)
        else:
            current_body.extend(table_lines)

    flush_current()

    chunks: list[ChunkRecord] = []
    ordinal = 1
    if intro_body:
        intro_chunk = _build_chunk(
            section_heading="Agreement overview",
            heading_path="Agreement overview",
            question_label="How should the rental agreement template be interpreted before the schedules are completed?",
            document_title_snapshot=title,
            body_text="\n".join(intro_body),
            source_locator="Agreement preface and master-copy instructions",
            semantic_split_decision="agreement preface block",
            oversized_section_handling="not needed",
        )
        chunks.append(_with_ordinal(intro_chunk, ordinal))
        ordinal += 1

    for heading, heading_path, body in sections:
        for chunk in _split_semantic_section(
            title=title,
            section_heading=heading,
            heading_path=heading_path,
            question_label=f"What does {heading.lower()} require?",
            paragraphs=body,
            base_source_locator=f"Heading path: {heading_path}",
            semantic_split_decision="heading block with ordered table support",
        ):
            chunks.append(_with_ordinal(chunk, ordinal))
            ordinal += 1
    return chunks


def _chunk_docx_numbered_sections(
    path: Path,
    title: str,
    include_top_level_numbers: tuple[int, ...],
    *,
    intro_heading: str | None,
    question_prefix: str,
) -> list[ChunkRecord]:
    document = Document(path)
    relevant_numbers = set(include_top_level_numbers)
    intro_body: list[str] = []
    current_top_level: str | None = None
    current_top_level_number: int | None = None
    current_subheading: str | None = None
    current_body: list[str] = []
    sections: list[tuple[str, str, list[str]]] = []

    def flush_current() -> None:
        nonlocal current_body
        if current_top_level and current_body:
            if current_subheading:
                section_heading = current_subheading
                heading_path = f"{current_top_level} > {current_subheading}"
            else:
                section_heading = current_top_level
                heading_path = current_top_level
            sections.append((section_heading, heading_path, current_body[:]))
        current_body = []

    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = " ".join(block.text.split())
            if not text:
                continue
            top_level_number = _extract_top_level_number(text)
            if top_level_number is not None:
                flush_current()
                current_top_level_number = top_level_number
                current_top_level = text if not relevant_numbers or top_level_number in relevant_numbers else None
                current_subheading = None
                continue

            if current_top_level_number not in relevant_numbers and relevant_numbers:
                if current_top_level_number is None:
                    intro_body.append(text)
                continue

            if re.match(r"^\d+[A-Z]\.\s", text):
                flush_current()
                current_subheading = text
                continue

            if current_top_level and _is_checklist_subheading(text):
                flush_current()
                current_subheading = text
                continue

            if current_top_level is None:
                intro_body.append(text)
            else:
                current_body.append(text)
            continue

        table_lines = _table_to_lines(block)
        if current_top_level_number not in relevant_numbers and relevant_numbers:
            if current_top_level_number is None:
                intro_body.extend(table_lines)
            continue
        if current_top_level is None:
            intro_body.extend(table_lines)
        else:
            current_body.extend(table_lines)

    flush_current()

    chunks: list[ChunkRecord] = []
    ordinal = 1
    if intro_heading and intro_body:
        intro_chunk = _build_chunk(
            section_heading=intro_heading,
            heading_path=intro_heading,
            question_label=f"{question_prefix} header be handled?".replace(" section", ""),
            document_title_snapshot=title,
            body_text="\n".join(intro_body),
            source_locator=f"{title} introduction",
            semantic_split_decision="document preface block",
            oversized_section_handling="not needed",
        )
        chunks.append(_with_ordinal(intro_chunk, ordinal))
        ordinal += 1

    for section_heading, heading_path, body in sections:
        lower_heading = _strip_numeric_prefix(section_heading).lower()
        question_label = f"{question_prefix} {lower_heading} be handled?"
        for chunk in _split_semantic_section(
            title=title,
            section_heading=section_heading,
            heading_path=heading_path,
            question_label=question_label,
            paragraphs=body,
            base_source_locator=f"Section: {heading_path}",
            semantic_split_decision="numbered section with ordered table support",
        ):
            chunks.append(_with_ordinal(chunk, ordinal))
            ordinal += 1
    return chunks


def _chunk_docx_linear_checklist_sections(path: Path, title: str) -> list[ChunkRecord]:
    document = Document(path)
    intro_body: list[str] = []
    current_heading: str | None = None
    current_body: list[str] = []
    sections: list[tuple[str, list[str]]] = []

    def flush_current() -> None:
        nonlocal current_body
        if current_heading and current_body:
            sections.append((current_heading, current_body[:]))
        current_body = []

    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = " ".join(block.text.split())
            if not text:
                continue
            if text == title:
                continue
            if text in {"Close-out", "Final notes"}:
                flush_current()
                current_heading = text
                continue
            if current_heading is None:
                intro_body.append(text)
            else:
                current_body.append(text)
            continue

        table_lines = _table_to_lines(block)
        if current_heading is None:
            intro_body.extend(table_lines)
        else:
            current_body.extend(table_lines)

    flush_current()

    chunks: list[ChunkRecord] = []
    ordinal = 1
    if intro_body:
        intro_chunk = _build_chunk(
            section_heading="Checklist overview",
            heading_path="Checklist overview",
            question_label="How should the rental close-out checklist be used?",
            document_title_snapshot=title,
            body_text="\n".join(intro_body),
            source_locator="Checklist overview",
            semantic_split_decision="checklist overview block",
            oversized_section_handling="not needed",
        )
        chunks.append(_with_ordinal(intro_chunk, ordinal))
        ordinal += 1

    for heading, body in sections:
        chunk = _build_chunk(
            section_heading=heading,
            heading_path=heading,
            question_label=f"What should be completed for {heading.lower()}?",
            document_title_snapshot=title,
            body_text="\n".join(body),
            source_locator=f"Section: {heading}",
            semantic_split_decision="linear checklist section",
            oversized_section_handling="not needed",
        )
        chunks.append(_with_ordinal(chunk, ordinal))
        ordinal += 1
    return chunks


def _chunk_governance_inventory(path: Path, title: str) -> list[ChunkRecord]:
    workbook = load_workbook(path, data_only=True, read_only=False)
    chunks: list[ChunkRecord] = []
    ordinal = 1
    ordinal = _append_sheet_intro_and_rows(
        workbook=workbook,
        chunks=chunks,
        ordinal=ordinal,
        title=title,
        sheet_name="Inventory",
        header_row=3,
        data_start_row=4,
        semantic_split_decision="inventory source-record row",
    )
    ordinal = _append_sheet_as_single_chunk(
        workbook=workbook,
        chunks=chunks,
        ordinal=ordinal,
        title=title,
        sheet_name="Overview",
        section_heading="Overview",
        question_label="How should the knowledge inventory overview be interpreted?",
        semantic_split_decision="governance overview sheet",
    )
    _append_sheet_as_single_chunk(
        workbook=workbook,
        chunks=chunks,
        ordinal=ordinal,
        title=title,
        sheet_name="Controlled Lists",
        section_heading="Controlled lists",
        question_label="What controlled list values support the knowledge inventory?",
        semantic_split_decision="governance controlled-list sheet",
    )
    return chunks


def _chunk_policy_decision_log(path: Path, title: str) -> list[ChunkRecord]:
    workbook = load_workbook(path, data_only=True, read_only=False)
    chunks: list[ChunkRecord] = []
    ordinal = 1
    ordinal = _append_sheet_intro_and_rows(
        workbook=workbook,
        chunks=chunks,
        ordinal=ordinal,
        title=title,
        sheet_name="Decision Log",
        header_row=4,
        data_start_row=5,
        semantic_split_decision="decision-log row",
    )
    ordinal = _append_sheet_intro_and_rows(
        workbook=workbook,
        chunks=chunks,
        ordinal=ordinal,
        title=title,
        sheet_name="Open Decisions",
        header_row=4,
        data_start_row=5,
        semantic_split_decision="open-decision row",
    )
    ordinal = _append_sheet_as_single_chunk(
        workbook=workbook,
        chunks=chunks,
        ordinal=ordinal,
        title=title,
        sheet_name="Overview",
        section_heading="Overview",
        question_label="How should the policy-governance overview be interpreted?",
        semantic_split_decision="policy overview sheet",
    )
    _append_sheet_as_single_chunk(
        workbook=workbook,
        chunks=chunks,
        ordinal=ordinal,
        title=title,
        sheet_name="Controlled Lists",
        section_heading="Controlled lists",
        question_label="What controlled list values support the policy decision log?",
        semantic_split_decision="policy controlled-list sheet",
    )
    return chunks


def _chunk_technical_inventory(path: Path, title: str) -> list[ChunkRecord]:
    workbook = load_workbook(path, data_only=True, read_only=False)
    chunks: list[ChunkRecord] = []
    ordinal = 1
    for sheet_name in (
        "Spaces",
        "Studio",
        "Retail-bar",
        "Conversation pit",
        "Waste Facilities  Waste Room",
        "1-1 Podcast Room",
        "Storage Room",
        "Back office",
        "Hallway & bathrooms",
        "Included equipment",
        "Technical capabilities",
    ):
        ordinal = _append_sheet_intro_and_rows(
            workbook=workbook,
            chunks=chunks,
            ordinal=ordinal,
            title=title,
            sheet_name=sheet_name,
            header_row=4,
            data_start_row=5,
            semantic_split_decision="technical-inventory row",
        )
    return chunks


def _chunk_capacity_space_rules(path: Path, title: str) -> list[ChunkRecord]:
    workbook = load_workbook(path, data_only=True, read_only=False)
    ws = workbook["Sheet1"]
    chunks: list[ChunkRecord] = []
    ordinal = 1

    intro_lines: list[str] = []
    for row_index in (1, 2):
        row_values = _row_values(ws, row_index)
        if row_values:
            intro_lines.append(" | ".join(row_values))
    if intro_lines:
        intro_chunk = _build_chunk(
            section_heading="Workbook overview",
            heading_path="Workbook overview",
            question_label="How should the capacity and space-use workbook be interpreted?",
            document_title_snapshot=title,
            body_text="\n".join(intro_lines),
            source_locator='Worksheet "Sheet1" overview',
            semantic_split_decision="capacity workbook overview",
            oversized_section_handling="not needed",
        )
        chunks.append(_with_ordinal(intro_chunk, ordinal))
        ordinal += 1

    current_section: str | None = None
    current_headers: list[str] = []
    for row_index in range(3, (ws.max_row or 0) + 1):
        values = _row_values(ws, row_index)
        if not values:
            continue
        first_value = values[0]
        if re.match(r"^\d+\.\s", first_value):
            current_section = first_value
            current_headers = []
            continue
        if current_section and not current_headers:
            current_headers = values
            continue
        if not current_section or not current_headers:
            continue
        if not re.match(r"^[A-Z]{2,}(?:-[A-Z]+)*-\d+", first_value):
            continue
        heading = _default_row_heading(values)
        chunk = _build_chunk(
            section_heading=heading,
            heading_path=f"{current_section} > {heading}",
            question_label=f"What does {heading.lower()} specify?",
            document_title_snapshot=title,
            body_text=_row_body_text(current_headers, values),
            source_locator=f'Worksheet "Sheet1", row {row_index}, record {first_value}',
            semantic_split_decision="capacity or access rule row",
            oversized_section_handling="not needed",
        )
        chunks.append(_with_ordinal(chunk, ordinal))
        ordinal += 1
    return chunks


def _chunk_catering_supplier_catalogue(path: Path, title: str) -> list[ChunkRecord]:
    workbook = load_workbook(path, data_only=True, read_only=False)
    chunks: list[ChunkRecord] = []
    ordinal = 1
    ordinal = _append_sheet_intro_and_rows(
        workbook=workbook,
        chunks=chunks,
        ordinal=ordinal,
        title=title,
        sheet_name="Catalogue",
        header_row=4,
        data_start_row=5,
        semantic_split_decision="catalogue row",
    )
    ordinal = _append_sheet_intro_and_rows(
        workbook=workbook,
        chunks=chunks,
        ordinal=ordinal,
        title=title,
        sheet_name="Catering & bar rules",
        header_row=4,
        data_start_row=5,
        semantic_split_decision="catering-rule row",
    )
    _append_sheet_as_single_chunk(
        workbook=workbook,
        chunks=chunks,
        ordinal=ordinal,
        title=title,
        sheet_name="Controlled lists",
        section_heading="Controlled lists",
        question_label="What controlled list values support the catering and supplier catalogue?",
        semantic_split_decision="catalogue controlled-list sheet",
    )
    return chunks


def _chunk_external_supplier_requirements(path: Path, title: str) -> list[ChunkRecord]:
    workbook = load_workbook(path, data_only=True, read_only=False)
    chunks: list[ChunkRecord] = []
    _append_sheet_intro_and_rows(
        workbook=workbook,
        chunks=chunks,
        ordinal=1,
        title=title,
        sheet_name="External supplier requirements",
        header_row=4,
        data_start_row=5,
        semantic_split_decision="external-supplier requirement row",
    )
    return chunks


def _append_sheet_intro_and_rows(
    *,
    workbook,
    chunks: list[ChunkRecord],
    ordinal: int,
    title: str,
    sheet_name: str,
    header_row: int,
    data_start_row: int,
    semantic_split_decision: str,
) -> int:
    ws = workbook[sheet_name]
    intro_lines: list[str] = []
    for row_index in range(1, header_row):
        row_values = _row_values(ws, row_index)
        if row_values:
            intro_lines.append(" | ".join(row_values))
    if intro_lines:
        intro_chunk = _build_chunk(
            section_heading=f"{sheet_name} overview",
            heading_path=f"{sheet_name} overview",
            question_label=f"What context governs the {sheet_name.lower()} sheet?",
            document_title_snapshot=title,
            body_text="\n".join(intro_lines),
            source_locator=f'Worksheet "{sheet_name}" overview',
            semantic_split_decision=f"{semantic_split_decision} overview",
            oversized_section_handling="not needed",
        )
        chunks.append(_with_ordinal(intro_chunk, ordinal))
        ordinal += 1

    headers = _row_values(ws, header_row)
    for row_index in range(data_start_row, (ws.max_row or 0) + 1):
        values = _row_values(ws, row_index)
        if not values:
            continue
        if values == headers:
            continue
        heading = _default_row_heading(values)
        identifier = values[0]
        chunk = _build_chunk(
            section_heading=heading,
            heading_path=f'{sheet_name} > {heading}',
            question_label=f"What does {heading.lower()} cover?",
            document_title_snapshot=title,
            body_text=_row_body_text(headers, values),
            source_locator=f'Worksheet "{sheet_name}", row {row_index}, record {identifier}',
            semantic_split_decision=semantic_split_decision,
            oversized_section_handling="not needed",
        )
        chunks.append(_with_ordinal(chunk, ordinal))
        ordinal += 1
    return ordinal


def _append_sheet_as_single_chunk(
    *,
    workbook,
    chunks: list[ChunkRecord],
    ordinal: int,
    title: str,
    sheet_name: str,
    section_heading: str,
    question_label: str,
    semantic_split_decision: str,
) -> int:
    ws = workbook[sheet_name]
    lines: list[str] = []
    for row_index in range(1, (ws.max_row or 0) + 1):
        row_values = _row_values(ws, row_index)
        if row_values:
            lines.append(f"Row {row_index}: {' | '.join(row_values)}")
    if not lines:
        return ordinal
    chunk = _build_chunk(
        section_heading=section_heading,
        heading_path=section_heading,
        question_label=question_label,
        document_title_snapshot=title,
        body_text="\n".join(lines),
        source_locator=f'Worksheet "{sheet_name}"',
        semantic_split_decision=semantic_split_decision,
        oversized_section_handling="not needed",
    )
    chunks.append(_with_ordinal(chunk, ordinal))
    return ordinal + 1


def _row_values(ws, row_index: int) -> list[str]:
    row = next(ws.iter_rows(min_row=row_index, max_row=row_index, values_only=True))
    values = [_cell_text(value) for value in row]
    while values and not values[-1]:
        values.pop()
    if not any(values):
        return []
    return values


def _default_row_heading(values: list[str]) -> str:
    nonempty = [value for value in values if value]
    if len(nonempty) >= 2:
        return f"{nonempty[0]} — {nonempty[1]}"
    return nonempty[0]


def _row_body_text(headers: list[str], values: list[str]) -> str:
    lines = [
        f"{header}: {value}"
        for header, value in zip(headers, values)
        if header and value
    ]
    return "\n".join(lines)
